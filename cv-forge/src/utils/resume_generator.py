"""
Resume generator module for creating PDF and Word resumes
Enhanced with multiple templates and improved formatting
"""
import os
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, Optional


class ResumeGenerator:
    """Generate professional resumes in PDF and Word formats with multiple templates"""

    # Template color schemes
    TEMPLATES = {
        'classic': {
            'name': 'Classic Professional',
            'primary_color': RGBColor(31, 71, 136),  # Navy blue
            'secondary_color': RGBColor(100, 100, 100),  # Gray
            'heading_font_size': 14,
            'body_font_size': 10,
            'name_font_size': 20,
            'use_lines': True,
            'use_icons': False,
        },
        'modern': {
            'name': 'Modern Clean',
            'primary_color': RGBColor(44, 62, 80),  # Dark slate
            'secondary_color': RGBColor(52, 152, 219),  # Bright blue
            'heading_font_size': 12,
            'body_font_size': 9,
            'name_font_size': 24,
            'use_lines': True,
            'use_icons': False,
        },
        'minimal': {
            'name': 'Minimal Elegant',
            'primary_color': RGBColor(0, 0, 0),  # Black
            'secondary_color': RGBColor(128, 128, 128),  # Gray
            'heading_font_size': 11,
            'body_font_size': 10,
            'name_font_size': 18,
            'use_lines': False,
            'use_icons': False,
        },
    }

    def __init__(self, output_dir: str = "output"):
        """
        Initialize resume generator

        Args:
            output_dir: Directory to save generated resumes
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def _get_template_config(self, template_name: str) -> Dict:
        """Get template configuration"""
        return self.TEMPLATES.get(template_name, self.TEMPLATES['classic'])

    def generate_pdf(self, resume_data: Dict, filename: Optional[str] = None, 
                     template: str = 'classic') -> str:
        """
        Generate PDF resume

        Args:
            resume_data: Dictionary containing resume information
            filename: Output filename (full path or just name)
            template: Template name ('classic', 'modern', 'minimal')

        Returns:
            Path to generated PDF
        """
        # Handle filename
        if filename is None:
            filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Ensure .pdf extension
        if not filename.endswith('.pdf'):
            filepath = filename if filename.endswith('.pdf') else f"{filename}.pdf"
        else:
            filepath = filename
            
        # Ensure output directory
        output_dir = os.path.dirname(filepath)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        else:
            filepath = os.path.join(self.output_dir, filepath)

        # Get template config
        config = self._get_template_config(template)

        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=0.5*inch,
            leftMargin=0.5*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )

        styles = getSampleStyleSheet()

        # Custom styles based on template
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=config['name_font_size'],
            textColor=config['primary_color'],
            spaceAfter=3,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

        section_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=config['heading_font_size'],
            textColor=config['primary_color'],
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            textTransform='uppercase',
            letterSpacing=2
        )

        body_style = ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontSize=config['body_font_size'],
            textColor=config['secondary_color'],
            alignment=TA_JUSTIFY,
            leading=14
        )

        contact_style = ParagraphStyle(
            'ContactInfo',
            parent=styles['Normal'],
            fontSize=9,
            textColor=config['secondary_color'],
            alignment=TA_CENTER,
            spaceAfter=6
        )

        # Separator line style (using underscores as line)
        separator_style = ParagraphStyle(
            'Separator',
            parent=styles['Normal'],
            fontSize=8,
            textColor=config['secondary_color'],
            alignment=TA_CENTER,
            spaceAfter=3,
            spaceBefore=3
        )

        content = []

        # Title (Name)
        name = resume_data.get('name', 'Your Name')
        title = Paragraph(name.upper(), title_style)
        content.append(title)

        # Professional Title
        if resume_data.get('title'):
            title_para = Paragraph(resume_data['title'], contact_style)
            content.append(title_para)

        # Contact Info
        if resume_data.get('contact'):
            contact_info = resume_data['contact'].replace('\n', ' | ')
            contact_para = Paragraph(contact_info, contact_style)
            content.append(contact_para)
            content.append(Spacer(1, 0.15*inch))
            
            # Add line separator
            if config['use_lines']:
                content.append(Paragraph("─" * 100, separator_style))
                content.append(Spacer(1, 0.15*inch))

        # Summary
        if resume_data.get('summary'):
            content.append(Paragraph("Professional Summary", section_style))
            if config['use_lines']:
                content.append(Paragraph("─" * 30, separator_style))
                content.append(Spacer(1, 0.1*inch))
            summary = Paragraph(resume_data['summary'], body_style)
            content.append(summary)
            content.append(Spacer(1, 0.15*inch))

        # Skills
        if resume_data.get('skills'):
            content.append(Paragraph("Skills", section_style))
            if config['use_lines']:
                content.append(Paragraph("─" * 15, separator_style))
                content.append(Spacer(1, 0.1*inch))
            skills = Paragraph(resume_data['skills'], body_style)
            content.append(skills)
            content.append(Spacer(1, 0.15*inch))

        # Experience
        if resume_data.get('experience'):
            content.append(Paragraph("Professional Experience", section_style))
            if config['use_lines']:
                content.append(Paragraph("─" * 40, separator_style))
                content.append(Spacer(1, 0.1*inch))
            experience = Paragraph(resume_data['experience'], body_style)
            content.append(experience)
            content.append(Spacer(1, 0.15*inch))

        # Education
        if resume_data.get('education'):
            content.append(Paragraph("Education", section_style))
            if config['use_lines']:
                content.append(Paragraph("─" * 20, separator_style))
                content.append(Spacer(1, 0.1*inch))
            education = Paragraph(resume_data['education'], body_style)
            content.append(education)
            content.append(Spacer(1, 0.15*inch))

        # Certifications
        if resume_data.get('certifications'):
            content.append(Paragraph("Certifications", section_style))
            if config['use_lines']:
                content.append(Paragraph("─" * 25, separator_style))
                content.append(Spacer(1, 0.1*inch))
            certifications = Paragraph(resume_data['certifications'], body_style)
            content.append(certifications)
            content.append(Spacer(1, 0.15*inch))

        # Projects
        if resume_data.get('projects'):
            content.append(Paragraph("Projects", section_style))
            if config['use_lines']:
                content.append(Paragraph("─" * 15, separator_style))
                content.append(Spacer(1, 0.1*inch))
            projects = Paragraph(resume_data['projects'], body_style)
            content.append(projects)
            content.append(Spacer(1, 0.15*inch))

        # Languages
        if resume_data.get('languages'):
            content.append(Paragraph("Languages", section_style))
            if config['use_lines']:
                content.append(Paragraph("─" * 15, separator_style))
                content.append(Spacer(1, 0.1*inch))
            languages = Paragraph(resume_data['languages'], body_style)
            content.append(languages)

        # References
        if resume_data.get('references'):
            content.append(Paragraph("References", section_style))
            if config['use_lines']:
                content.append(Paragraph("─" * 15, separator_style))
                content.append(Spacer(1, 0.1*inch))
            references = Paragraph(resume_data['references'], body_style)
            content.append(references)

        doc.build(content)
        return filepath

    def generate_docx(self, resume_data: Dict, filename: Optional[str] = None,
                      template: str = 'classic') -> str:
        """
        Generate Word document resume

        Args:
            resume_data: Dictionary containing resume information
            filename: Output filename (full path or just name)
            template: Template name ('classic', 'modern', 'minimal')

        Returns:
            Path to generated Word document
        """
        # Handle filename
        if filename is None:
            filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Ensure .docx extension
        if not filename.endswith('.docx'):
            filepath = filename if filename.endswith('.docx') else f"{filename}.docx"
        else:
            filepath = filename
            
        # Ensure output directory
        output_dir = os.path.dirname(filepath)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        else:
            filepath = os.path.join(self.output_dir, filepath)

        # Get template config
        config = self._get_template_config(template)

        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

        # Title (Name)
        name = resume_data.get('name', 'Your Name')
        title = doc.add_paragraph(name.upper())
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format = title.runs[0]
        title_format.font.size = Pt(config['name_font_size'])
        title_format.font.bold = True
        title_format.font.color.rgb = config['primary_color']
        title_format.font.name = 'Calibri'

        # Professional Title
        if resume_data.get('title'):
            title_para = doc.add_paragraph(resume_data['title'])
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = config['secondary_color']

        # Contact Info
        if resume_data.get('contact'):
            contact = doc.add_paragraph(resume_data['contact'].replace('\n', ' | '))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact.runs[0]
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = config['secondary_color']

        # Add separator line
        if config['use_lines']:
            doc.add_paragraph('_' * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        def add_section(heading: str, content_text: str, short_line: bool = False):
            """Helper to add a section"""
            if not content_text:
                return
            
            # Heading
            heading_para = doc.add_paragraph(heading.upper())
            heading_run = heading_para.runs[0]
            heading_run.font.bold = True
            heading_run.font.size = Pt(config['heading_font_size'])
            heading_run.font.color.rgb = config['primary_color']
            heading_run.font.name = 'Calibri'
            
            # Short line under heading
            if config['use_lines']:
                line_length = int(short_line * 20) if short_line else 30
                doc.add_paragraph('_' * line_length)
            
            # Content
            doc.add_paragraph(content_text)
            para = doc.paragraphs[-1]
            para.runs[0].font.size = Pt(config['body_font_size'])
            para.runs[0].font.color.rgb = config['secondary_color']
            para.runs[0].font.name = 'Calibri'
            
            doc.add_paragraph()

        # Summary
        add_section("Professional Summary", resume_data.get('summary', ''), short_line=2)

        # Skills
        add_section("Skills", resume_data.get('skills', ''), short_line=1)

        # Experience
        add_section("Professional Experience", resume_data.get('experience', ''), short_line=3)

        # Education
        add_section("Education", resume_data.get('education', ''), short_line=2)

        # Certifications
        add_section("Certifications", resume_data.get('certifications', ''), short_line=2)

        # Projects
        add_section("Projects", resume_data.get('projects', ''), short_line=1)

        # Languages
        add_section("Languages", resume_data.get('languages', ''), short_line=2)

        # References
        add_section("References", resume_data.get('references', ''), short_line=2)

        doc.save(filepath)
        return filepath

    def generate_all_formats(self, resume_data: Dict, filename: Optional[str] = None,
                             template: str = 'classic') -> Dict[str, str]:
        """
        Generate resume in all formats (PDF and Word)

        Args:
            resume_data: Dictionary containing resume information
            filename: Base filename (without extension)
            template: Template name

        Returns:
            Dictionary with paths to generated files
        """
        results = {}
        
        if filename is None:
            filename = f"resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Generate PDF
        pdf_path = self.generate_pdf(resume_data, filename, template)
        results['pdf'] = pdf_path
        
        # Generate Word
        docx_path = self.generate_docx(resume_data, filename, template)
        results['docx'] = docx_path
        
        return results

    def get_available_templates(self) -> Dict[str, str]:
        """
        Get available template names and descriptions

        Returns:
            Dictionary mapping template keys to display names
        """
        return {key: config['name'] for key, config in self.TEMPLATES.items()}
