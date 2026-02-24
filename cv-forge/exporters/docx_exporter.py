"""
DOCX Exporter for CV-Forge.
Generates ATS-friendly Word documents using python-docx.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))
from models.resume import Resume


class DOCXExporter:
    """Export Resume to ATS-friendly DOCX format."""
    
    def __init__(self, resume: Resume):
        self.resume = resume
        self.doc = Document()
        
        # Set up base styles for ATS compatibility
        self._setup_styles()
    
    def _setup_styles(self):
        """Configure document styles for ATS compatibility."""
        # ATS-friendly: simple fonts, no columns, no tables for layout
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
    
    def export(self, filepath: str) -> None:
        """
        Export the resume to a DOCX file.
        
        Args:
            filepath: Path to save the DOCX file
        """
        self._add_header()
        self._add_profile()
        self._add_education()
        self._add_certifications()
        self._add_experiences()
        self._add_skills()
        
        self.doc.save(filepath)
    
    def _add_header(self) -> None:
        """Add resume header with contact information."""
        # Name - centered, bold, larger
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(self.resume.full_name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        
        # Contact info - centered, simple text format
        contact_lines = []
        if self.resume.phone:
            contact_lines.append(self.resume.phone)
        if self.resume.email:
            contact_lines.append(self.resume.email)
        if self.resume.linkedin:
            contact_lines.append(self.resume.linkedin)
        if self.resume.address:
            contact_lines.append(self.resume.address)
        
        if contact_lines:
            contact_para = self.doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" | ".join(contact_lines))
            contact_run.font.size = Pt(10)
        
        # Add spacing after header
        self.doc.add_paragraph()
    
    def _add_profile(self) -> None:
        """Add profile/summary section."""
        if not self.resume.profile:
            return
        
        self._add_section_title("PROFIL")
        profile_para = self.doc.add_paragraph(self.resume.profile)
        profile_para.paragraph_format.space_after = Pt(12)
    
    def _add_education(self) -> None:
        """Add education section."""
        if not self.resume.education:
            return
        
        self._add_section_title("FORMATION")
        
        for edu in self.resume.education:
            # Format: Diploma – Institution – Dates
            edu_text = f"{edu.diploma} – {edu.institution} – {edu.dates}"
            edu_para = self.doc.add_paragraph(edu_text)
            edu_para.paragraph_format.space_after = Pt(6)
    
    def _add_certifications(self) -> None:
        """Add certifications section."""
        if not self.resume.certifications:
            return
        
        self._add_section_title("CERTIFICATION")
        
        for cert in self.resume.certifications:
            # Format: Name – Organization – Year
            cert_text = f"{cert.name} – {cert.organization} – {cert.year}"
            cert_para = self.doc.add_paragraph(cert_text)
            cert_para.paragraph_format.space_after = Pt(6)
    
    def _add_experiences(self) -> None:
        """Add experience section."""
        if not self.resume.experiences:
            return
        
        self._add_section_title("EXPERIENCES PROFESSIONNELLES")
        
        for exp in self.resume.experiences:
            # Format: Position – Company – City – Dates
            dates_range = f"{exp.start_date} – {exp.end_date}"
            exp_header = f"{exp.position} – {exp.company} – {exp.city} – {dates_range}"
            
            exp_para = self.doc.add_paragraph(exp_header)
            exp_para.paragraph_format.space_after = Pt(6)
            
            # Add bullet points
            for bullet in exp.bullets:
                bullet_para = self.doc.add_paragraph(style='List Bullet')
                bullet_para.add_run(bullet)
                bullet_para.paragraph_format.space_after = Pt(3)
    
    def _add_skills(self) -> None:
        """Add skills section (hard and soft skills)."""
        has_skills = self.resume.skills_hard or self.resume.skills_soft
        if not has_skills:
            return
        
        self._add_section_title("COMPETENCES")
        
        if self.resume.skills_hard:
            skills_para = self.doc.add_paragraph()
            skills_para.add_run("Compétences Techniques:\n").bold = True
            for skill in self.resume.skills_hard:
                skills_para.add_run(f"• {skill}\n")
            skills_para.paragraph_format.space_after = Pt(12)
        
        if self.resume.skills_soft:
            skills_para = self.doc.add_paragraph()
            skills_para.add_run("Compétences Comportementales:\n").bold = True
            for skill in self.resume.skills_soft:
                skills_para.add_run(f"• {skill}\n")
    
    def _add_section_title(self, title: str) -> None:
        """Add a section title with consistent formatting."""
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_run.all_caps = True
        title_para.paragraph_format.space_before = Pt(12)
        title_para.paragraph_format.space_after = Pt(6)
        
        # Add underline
        underline_para = self.doc.add_paragraph()
        underline_run = underline_para.add_run("─" * 50)
        underline_run.font.size = Pt(8)
        underline_para.paragraph_format.space_after = Pt(6)
